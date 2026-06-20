from __future__ import annotations

import base64
import io
import re
from functools import lru_cache
from dataclasses import dataclass
from typing import Any

import fitz  # PyMuPDF
from PIL import Image, ImageOps, ImageStat


def normalize_spaces(text: str) -> str:
    return " ".join(str(text or "").split()).strip()


def _clean_ocr_line(text: str) -> str:
    cleaned = normalize_spaces(text)
    if not cleaned:
        return ""

    tokens = cleaned.split()
    if len(tokens) < 2:
        return cleaned

    generic_prefixes = {"general", "ocr", "text", "image", "page"}

    for index, token in enumerate(tokens):
        token_core = token.strip(".,;:!?()[]{}<>\"'")
        if not token_core:
            continue
        if any(ch.isdigit() for ch in token_core):
            break
        if any(ch.isupper() for ch in token_core):
            if 0 < index <= 3:
                prefix = tokens[:index]
                if all(
                    prefix_token.strip(".,;:!?()[]{}<>\"'").isascii()
                    and prefix_token.strip(".,;:!?()[]{}<>\"'").islower()
                    for prefix_token in prefix
                ):
                    trimmed = normalize_spaces(" ".join(tokens[index:]))
                    return trimmed or cleaned
            break

    if cleaned.isascii():
        lowered = cleaned.lower()
        if lowered in generic_prefixes:
            return ""
        if len(tokens) <= 2 and all(token.islower() for token in tokens) and any(token in generic_prefixes for token in tokens):
            return ""

    return cleaned


def _prepare_ocr_image(image: Image.Image, *, max_side: int = 1400) -> Image.Image:
    width, height = image.size
    longest_side = max(width, height)
    if longest_side <= max_side or longest_side <= 0:
        return image

    scale = max_side / float(longest_side)
    new_size = (max(1, round(width * scale)), max(1, round(height * scale)))
    resample = getattr(Image, "Resampling", Image).LANCZOS
    return image.resize(new_size, resample)


def _ocr_image_with_paddle_image(image: Image.Image, *, ocr_langs: str = "vie+eng") -> list[str]:
    import numpy as np

    image = ImageOps.exif_transpose(image).convert("RGB")
    image = _prepare_ocr_image(image)
    lang = _resolve_paddleocr_lang(ocr_langs)
    engine = _get_paddleocr_engine(lang)
    image_array = np.asarray(image)
    try:
        result = engine.ocr(image_array, cls=True)
    except TypeError:
        result = engine.ocr(image_array)
    return _extract_paddleocr_lines(result)


def is_pdf(file_name: str, mime_type: str, data: bytes) -> bool:
    lowered_name = (file_name or "").lower()
    lowered_mime = (mime_type or "").lower()
    if lowered_mime == "application/pdf" or lowered_name.endswith(".pdf"):
        return True
    return data[:4] == b"%PDF"


def is_text_document(file_name: str, mime_type: str) -> bool:
    lowered_name = (file_name or "").lower()
    lowered_mime = (mime_type or "").lower()
    return lowered_mime in {
        "text/plain",
        "text/markdown",
        "application/json",
        "application/xml",
        "text/csv",
    } or lowered_name.endswith((".txt", ".md", ".csv", ".json", ".xml"))


def is_docx(file_name: str, mime_type: str) -> bool:
    lowered_name = (file_name or "").lower()
    lowered_mime = (mime_type or "").lower()
    return lowered_mime in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    } or lowered_name.endswith(".docx")


def is_image(file_name: str, mime_type: str) -> bool:
    lowered_name = (file_name or "").lower()
    lowered_mime = (mime_type or "").lower()
    if lowered_mime.startswith("image/"):
        return True
    return lowered_name.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff"))


def pdf_has_embedded_text(data: bytes) -> bool:
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception:
        return False
    try:
        for page in doc:
            if normalize_spaces(page.get_text("text")):
                return True
        return False
    finally:
        doc.close()


def _resolve_paddleocr_lang(ocr_langs: str) -> str:
    normalized = (ocr_langs or "").lower()
    if any(token in normalized for token in ("vi", "vie", "vietnam")):
        return "vi"
    if any(token in normalized for token in ("en", "eng", "english")):
        return "en"
    return "en"


@lru_cache(maxsize=4)
def _get_paddleocr_engine(lang: str):
    try:
        from paddleocr import PaddleOCR
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError("paddleocr is not installed.") from exc

    common_kwargs = {
        "lang": lang,
        "device": "cpu",
        "use_doc_orientation_classify": False,
        "use_doc_unwarping": False,
        "use_textline_orientation": False,
        "enable_mkldnn": False,
        "cpu_threads": 1,
    }
    try:
        return PaddleOCR(**common_kwargs)
    except TypeError:
        try:
            return PaddleOCR(lang=lang)
        except TypeError:
            return PaddleOCR(use_angle_cls=True, lang=lang)


def _extract_paddleocr_lines(result: Any) -> list[str]:
    lines: list[str] = []
    seen: set[str] = set()

    def add_line(text: str, *, score: float | None = None) -> None:
        if score is not None and score < 0.25:
            return
        cleaned = _clean_ocr_line(text)
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            lines.append(cleaned)

    def walk(node: Any) -> None:
        if node is None:
            return
        if isinstance(node, str):
            add_line(node)
            return
        if isinstance(node, dict):
            rec_texts = node.get("rec_texts")
            if isinstance(rec_texts, (list, tuple)):
                rec_scores = node.get("rec_scores")
                if not isinstance(rec_scores, (list, tuple)):
                    rec_scores = []
                for index, item in enumerate(rec_texts):
                    score = rec_scores[index] if index < len(rec_scores) and isinstance(rec_scores[index], (int, float)) else None
                    if isinstance(item, str):
                        add_line(item, score=score)
                if rec_texts:
                    return
            text = node.get("text")
            if isinstance(text, str):
                add_line(text)
                return
            for value in node.values():
                if isinstance(value, (str, dict, list, tuple)):
                    walk(value)
            return
        if isinstance(node, (list, tuple)):
            if len(node) == 2 and isinstance(node[1], (list, tuple)) and node[1]:
                text_part = node[1][0]
                if isinstance(text_part, str):
                    add_line(text_part)
                    return
            for item in node:
                walk(item)

    walk(result)
    return lines


def _ocr_image_with_paddle(data: bytes, *, ocr_langs: str = "vie+eng") -> list[str]:
    image = Image.open(io.BytesIO(data))
    return _ocr_image_with_paddle_image(image, ocr_langs=ocr_langs)


def extract_text_from_text_file(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1258", "latin-1"):
        try:
            text = data.decode(encoding)
            return normalize_spaces(text)
        except UnicodeDecodeError:
            continue
    return ""


def extract_text_from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError("python-docx is not installed.") from exc

    with io.BytesIO(data) as buffer:
        doc = Document(buffer)
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        tables: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    tables.append(row_text)
        return normalize_spaces("\n".join(paragraphs + tables))


def extract_text_from_pdf(data: bytes, *, ocr_langs: str = "vie+eng", ocr_if_empty: bool = True) -> tuple[str, dict[str, Any]]:
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        page_texts: list[str] = []
        page_refs: list[dict[str, Any]] = []
        ocr_pages = 0
        for page_index, page in enumerate(doc, start=1):
            text = normalize_spaces(page.get_text("text"))
            if not text and ocr_if_empty:
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                text = normalize_spaces("\n".join(_ocr_image_with_paddle_image(image, ocr_langs=ocr_langs)))
                if text:
                    ocr_pages += 1
            page_texts.append(text)
            if text:
                page_refs.append(
                    {
                        "page": page_index,
                        "snippet": normalize_spaces(text)[:240],
                    }
                )
        cleaned_pages = [normalize_spaces(part) for part in page_texts if normalize_spaces(part)]
        merged = "\n\n".join(cleaned_pages)
        return merged, {
            "pages": doc.page_count,
            "ocr_pages": ocr_pages,
            "has_embedded_text": bool(merged),
            "page_texts": cleaned_pages,
            "page_refs": page_refs[:12],
        }
    finally:
        doc.close()


def image_to_bytes(image: Image.Image, format: str = "PNG") -> bytes:
    with io.BytesIO() as buffer:
        image.save(buffer, format=format)
        return buffer.getvalue()


def extract_text_from_image_bytes(data: bytes, *, ocr_langs: str = "vie+eng") -> str:
    try:
        lines = _ocr_image_with_paddle(data, ocr_langs=ocr_langs)
    except Exception:
        return ""
    return normalize_spaces("\n".join(lines))


def image_basic_stats(data: bytes) -> dict[str, Any]:
    image = Image.open(io.BytesIO(data))
    image = ImageOps.exif_transpose(image).convert("RGB")
    width, height = image.size
    stat = ImageStat.Stat(image)
    mean_brightness = sum(stat.mean) / max(len(stat.mean), 1)
    return {
        "width": width,
        "height": height,
        "mode": image.mode,
        "aspect_ratio": round(width / height, 4) if height else None,
        "brightness": round(mean_brightness, 2),
        "estimated_quality": "good" if mean_brightness > 60 else "low_light_or_dark",
    }


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    meta: dict[str, Any]


def extract_document_text(
    data: bytes,
    *,
    file_name: str = "",
    mime_type: str = "",
    ocr_langs: str = "vie+eng",
) -> ExtractedDocument:
    if is_text_document(file_name, mime_type):
        return ExtractedDocument(text=extract_text_from_text_file(data), meta={"kind": "text"})
    if is_docx(file_name, mime_type):
        return ExtractedDocument(text=extract_text_from_docx(data), meta={"kind": "docx"})
    if is_pdf(file_name, mime_type, data):
        text, meta = extract_text_from_pdf(data, ocr_langs=ocr_langs)
        return ExtractedDocument(text=text, meta={"kind": "pdf", **meta})

    text = extract_text_from_text_file(data)
    return ExtractedDocument(text=text, meta={"kind": "plain_text"})


def parse_structured_fields(text: str) -> dict[str, Any]:
    normalized = text or ""
    emails = list(dict.fromkeys(re.findall(r"[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}", normalized, flags=re.IGNORECASE)))
    phones = list(dict.fromkeys(re.findall(r"(?:\+?\d[\d\s().-]{7,}\d)", normalized)))
    dates = list(dict.fromkeys(re.findall(r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})\b", normalized)))
    money = list(dict.fromkeys(re.findall(r"\b\d[\d.,]*(?:\s?(?:vnd|đ|dollar|usd|usd))?\b", normalized, flags=re.IGNORECASE)))
    urls = list(dict.fromkeys(re.findall(r"https?://[^\s)>\"']+", normalized, flags=re.IGNORECASE)))
    return {
        "emails": emails[:10],
        "phones": phones[:10],
        "dates": dates[:10],
        "money": money[:10],
        "urls": urls[:10],
    }


def compact_text(text: str, limit: int = 12000) -> str:
    cleaned = normalize_spaces(text)
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rstrip() + "..."


def image_to_data_url(data: bytes, mime_type: str) -> str:
    mime = mime_type or "image/png"
    encoded = base64.b64encode(data).decode("ascii")
    return f"data:{mime};base64,{encoded}"
